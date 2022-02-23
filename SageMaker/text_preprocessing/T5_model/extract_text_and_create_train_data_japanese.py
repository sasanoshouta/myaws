import boto3
import pandas as pd
import numpy as np
import gc
from janome.tokenizer import Tokenizer
import docx

# wordデータをpython-docxで読み込み、文章、表情報を取得
doc = docx.Document('japanese.docx')
para = doc.paragraphs
para_list = [para.text.replace("\t", " ") for para in doc.paragraphs if (para.text != '') and (para.text != '\n') and (para.text != '\t')]
# 基データ読み込み
new_sample = pd.read_excel('./.sample_ja.xlsx', engine='openpyxl', sheet_name='sheet1', header=9).reset_index(drop=True)

# 日本語文章を品詞分解し、助詞で区切り名詞区、節を作成
tmp = list()
words = list()
hinshi = ['名詞', '接続詞', '動詞', '接頭詞']
finish = '記号'

t = Tokenizer()
for i in range(len(new_sample)):
    input_data = new_sample.loc[i, 'input']
    for j in input_data:
        for token in t.tokenize(j):
            if token.part_of_speech.split(',')[0] in hinshi:
                tmp.append(token.surface)
            elif token.part_of_speech.split(',')[1] == '空白':
                tmp.append(token.surface)
            elif (token.part_of_speech.split(',')[0] == '助詞') or (token.part_of_speech.split(',')[1] == '句点') or (token.part_of_speech.split(',')[1] == '読点'):
                words.append(''.join(tmp))
                tmp = list()
            else:
                pass
        words.append(''.join(tmp))
        tmp = list()
    words = [s for s in words if len(s) != 0]
    new_sample.loc[i, 'input'] = list(set(words))
    words = list()

# 不要な記号の削除
code_regex = re.compile('[\t\s!"#$%&\'\\\\()*+,-./:;；：<=>?@[\\]^_`{|}~○｢｣「」〔〕“”〈〉'\
                '『』【】＆＊（）＄＃＠？！｀＋￥¥％♪…◇→←↓↑｡･ω･｡ﾟ´∀｀ΣДｘ⑥◎©︎♡★☆▽※ゞノ〆εσ＞＜┌┘]・・・・・・・・・・・・・・・')

# python-docxで抽出したテキストデータから学習用データを作成
for i in range(len(new_sample)):
    check = sample['input'][i]
    tmp_text = list()
    for c_text in check:
        for page in para_list:
            if c_text in page:
                page = page.replace('\n', '、')
                page = code_regex.sub(' ', page)
                tmp_text.append(page)
            else:
                pass
    new_sample.loc[i, 'input'] = tmp_text
    
# 日→英に翻訳してdfに格納
# AWS Translateにて
translate = boto3.client('translate')
index = list(new_sample.index)
tmp = list()
new_sample['input_en'] = '[]'
new_sample['確認事項_en'] = '[]'

for ind in range(len(new_sample)):
    for txt in new_sample.loc[ind, 'input']:
        result_input = translate.translate_text(Text=txt, SourceLanguageCode='ja', TargetLanguageCode='en')
        tmp.append(result_input['TranslatedText'])
    result_testitem = translate.translate_text(Text=new_sample.loc[ind, '確認事項'], SourceLanguageCode='ja', TargetLanguageCode='en')
    new_sample.loc[ind, 'input_en'] = " ".join(tmp)
    new_sample.loc[ind, '確認事項_en'] = result_testitem['TranslatedText']
    tmp = list()
    
new_sample.to_csv('train_ja.csv')
